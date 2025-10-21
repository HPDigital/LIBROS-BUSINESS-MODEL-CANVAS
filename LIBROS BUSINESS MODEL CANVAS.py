"""
LIBROS BUSINESS MODEL CANVAS
"""

#!/usr/bin/env python
# coding: utf-8

# In[2]:


import os
from docx import Document
from dotenv import load_dotenv
from langchain.prompts import PromptTemplate
from langchain.vectorstores import FAISS
import nltk
from nltk.tokenize import sent_tokenize
from pydantic import BaseModel
from langchain_openai import ChatOpenAI, OpenAIEmbeddings

# Descargar recursos de NLTK (descomentar si es la primera vez)
# nltk.download('punkt')

# =============================================================================
# CONFIGURACIÓN Y CLIENTES
# =============================================================================

load_dotenv()
API_KEY = os.getenv("OPENAI_API_KEY")
if not API_KEY:
    raise ValueError("La clave de la API de OpenAI no se encontró en las variables de entorno.")

# Inicializar cliente LLM (se configura el modelo y otros parámetros según se requiera)
llm = ChatOpenAI(api_key=API_KEY, model="o3-mini-2025-01-31", temperature=1)

# Inicializar embeddings para vectorización
embeddings = OpenAIEmbeddings(api_key=API_KEY)

# =============================================================================
# FUNCIONES PARA EXTRAER Y PROCESAR TEXTO DE ARCHIVOS PDF
# =============================================================================

def get_pdf_text(ruta_pdf: str, max_tokens: int = 100000) -> str:
    """
    Lee y extrae el texto de un PDF, limitándolo a max_tokens tokens.
    """
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(ruta_pdf)
        tokens_totales = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                tokens_page = page_text.split()
                tokens_totales.extend(tokens_page)
                if len(tokens_totales) >= max_tokens:
                    tokens_totales = tokens_totales[:max_tokens]
                    break
        print(f"Total tokens leídos del PDF: {len(tokens_totales)}")
        return " ".join(tokens_totales)
    except Exception as e:
        print(f"Error al leer el PDF: {e}")
        return ""

def chunk_text_to_sentences(texto: str, min_words: int = 70) -> list:
    """
    Divide el texto en oraciones y filtra aquellas que tengan más de min_words palabras.
    """
    oraciones = sent_tokenize(texto, language='spanish')
    return [oracion.strip() for oracion in oraciones if len(oracion.split()) > min_words]

# =============================================================================
# CREACIÓN DE VECTORSTORES (EMBEDDINGS) BASADOS EN PDF
# =============================================================================

def create_style_vectorstore(ruta_estilo: str) -> FAISS:
    """
    Crea un vectorstore a partir de un archivo PDF de estilo.
    """
    if not ruta_estilo.lower().endswith('.pdf'):
        raise ValueError("El archivo de estilo debe estar en formato PDF.")

    texto_estilo = get_pdf_text(ruta_estilo, max_tokens=100000)
    oraciones = chunk_text_to_sentences(texto_estilo, min_words=70)
    if not oraciones:
        raise ValueError("No se encontraron oraciones válidas en el archivo de estilo.")
    return FAISS.from_texts(oraciones, embeddings)

def create_respuestas_vectorstore(ruta_respuestas: str) -> FAISS:
    """
    Crea un vectorstore a partir de un archivo PDF con frases de respuestas de otros libros.
    """
    if not ruta_respuestas.lower().endswith('.pdf'):
        raise ValueError("El archivo de respuestas debe estar en formato PDF.")

    texto_respuestas = get_pdf_text(ruta_respuestas, max_tokens=100000)
    oraciones = chunk_text_to_sentences(texto_respuestas, min_words=40)
    if not oraciones:
        raise ValueError("No se encontraron oraciones válidas en el archivo de respuestas.")
    return FAISS.from_texts(oraciones, embeddings)

def retrieve_style_text(vectorstore: FAISS, query: str = "Extrae las oraciones que mejor representen un estilo formal, académico y fluido", k: int = 10) -> str:
    """
    Recupera oraciones representativas del estilo deseado.
    """
    resultados = vectorstore.similarity_search(query, k=k)
    oraciones = [doc.page_content for doc in resultados]
    return "\n".join(oraciones)

def retrieve_respuestas_context(vectorstore: FAISS, query: str, k: int = 5) -> str:
    """
    Recupera frases relevantes del vectorstore de respuestas de otros libros.
    """
    resultados = vectorstore.similarity_search(query, k=k)
    frases = [doc.page_content for doc in resultados]
    return "\n".join(frases)

# =============================================================================
# TEMPLATES PARA GENERACIÓN DE RESPUESTAS
# =============================================================================

PROMPT_SUBSUBTITULO_EXPLICACION = PromptTemplate(
    input_variables=["titulo", "subtitulo", "subsubtitulo", "contexto_libros"],
    template="""Proporciona una explicación clara y detallada con dos ejemplos ilustrativos para el subtítulo de tercer nivel '{subsubtitulo}', que forma parte del subtítulo '{subtitulo}' del título '{titulo}' en el contexto del Business Model Canvas.
Utiliza el siguiente contexto extraído de fuentes académicas y otros libros:
{contexto_libros}
Asegúrate de que la respuesta incluya una exposición teórica detallada, profunda y académica, analice el concepto y su utilidad y presente dos ejemplos prácticos en un escenario real.
La redacción debe ser clara, debe ser formal y debe ser académica, con oraciones extensas (al menos 50 palabras) y al menos tres oraciones concetadas por un punto seguido.
Respuesta:"""
)

PROMPT_APPLY_HUMANO = PromptTemplate(
    input_variables=["contenido_original", "estilo"],
    template="""Ajusta tu forma de redacción para cumplir con los siguientes criterios, asegurando un texto que se asemeje a una respuesta completamente humana y alejada de los patrones de escritura típicos de inteligencia artificial.

Estilo de Redacción:
Imita la fluidez, la coherencia y la sofisticación de un texto académico extraído de un documento formal en español. Para ello, ten en cuenta las siguientes directrices:

Extensión y Complejidad de las Frases:
Redacta oraciones extensas que integren varias ideas de manera armónica, utilizando conectores fluidos en español.

Estructura Cohesiva y Fragmentación Adecuada:
Organiza el contenido en un mínimo de dos y un máximo de cinco oraciones interconectadas, asegurando una redacción natural y evitando frases inconexas.

Claridad y Consistencia:
Cada oración debe mantener un balance entre precisión y fluidez, evitando repeticiones innecesarias o explicaciones redundantes.

Uso Estratégico de Conectores:
Minimiza los puntos aislados y favorece una conexión fluida de ideas mediante conjunciones variadas, sin caer en estructuras mecánicas.

Diversificación Sintáctica:
Utiliza una combinación de oraciones subordinadas, construcciones con participios y elementos discursivos que den variedad y riqueza al texto.

Tono Académico pero Accesible:
Mantén un estilo formal y estructurado, pero sin tecnicismos excesivos que dificulten la comprensión. Piensa en la redacción de un libro universitario.

Extensión de las Frases:
Cada frase debe contener más de 50 palabras, integrando al menos tres proposiciones, asegurando un flujo de información profundo y detallado.

Instrucciones Finales:
Redacta el siguiente contenido aplicando todas las reglas anteriores, asegurando que el resultado sea indistinguible de un texto humano de alto nivel académico. Evita frases cortas, patrones predecibles o un tono excesivamente robótico. Usa la referencia proporcionada a continuación como base para emular el estilo lingüístico y discursivo.

Ejemplos de redacción: Utiliza de ejemplo de redacción las frases siguientes {estilo}

Reformula el siguiente contenido:

{contenido_original}
aplicándole el siguiente estilo humano deseado:
{estilo}
El resultado debe ser una redacción académica, fluida, con oraciones extensas (más de 50 palabras) y al menos tres oraciones , separadas por un punto seguido, manteniendo un tono cercano pero profesional.
Respuesta:"""
)

# =============================================================================
# FUNCIONES DE GENERACIÓN DE TEXTO
# =============================================================================

def generate_subsubtitulo_explanation(titulo: str, subtitulo: str, subsubtitulo: str, vectorstore_respuestas: FAISS, k: int = 10) -> str:
    """
    Genera la explicación y el ejemplo para un subtítulo de tercer nivel,
    integrando contexto extraído mediante RAG de otros libros.
    """
    query = f"Explicación y ejemplo académico para '{subsubtitulo}' en el contexto de '{titulo}' > '{subtitulo}'"
    contexto_libros = retrieve_respuestas_context(vectorstore_respuestas, query, k)
    formatted_prompt = PROMPT_SUBSUBTITULO_EXPLICACION.format(
        titulo=titulo,
        subtitulo=subtitulo,
        subsubtitulo=subsubtitulo,
        contexto_libros=contexto_libros
    )
    response = llm.invoke(formatted_prompt)
    return response.content.strip()

def apply_human_style_to_text(contenido_original: str, estilo: str) -> str:
    """
    Transforma el contenido original aplicándole el estilo humano deseado.
    """
    formatted_prompt = PROMPT_APPLY_HUMANO.format(
        contenido_original=contenido_original,
        estilo=estilo
    )
    response = llm.invoke(formatted_prompt)
    return response.content.strip()

# =============================================================================
# GENERACIÓN DEL DOCUMENTO WORD A PARTIR DE UNA ESTRUCTURA JSON
# =============================================================================

def generate_word_document(json_data: dict, output_path: str, ruta_estilo: str, ruta_respuestas: str):
    """
    Procesa un JSON con la estructura jerárquica del Business Model Canvas y genera,
    para cada subtítulo de tercer nivel, una explicación académica detallada enriquecida
    con contexto de otros libros (usando RAG) y adaptada al estilo humano deseado.
    """
    # Crear vectorstores
    vectorstore_estilo = create_style_vectorstore(ruta_estilo)
    texto_estilo = retrieve_style_text(vectorstore_estilo)
    vectorstore_respuestas = create_respuestas_vectorstore(ruta_respuestas)

    doc = Document()
    doc.add_heading("Libro Business Model Canvas", level=1)

    for titulo, subtitulos in json_data.items():
        doc.add_heading(titulo, level=1)
        for subtitulo, subsubtitulos in subtitulos.items():
            doc.add_heading(subtitulo, level=2)
            for subsubtitulo in subsubtitulos:
                print(f"Procesando: {titulo} > {subtitulo} > {subsubtitulo}")
                # Generar explicación con contexto relevante
                explicacion = generate_subsubtitulo_explanation(titulo, subtitulo, subsubtitulo, vectorstore_respuestas, k=5)
                # Aplicar estilo humano deseado a la explicación
                explicacion_estilizada = apply_human_style_to_text(explicacion, texto_estilo)
                doc.add_heading(subsubtitulo, level=3)
                doc.add_paragraph(explicacion_estilizada)

    doc.save(output_path)
    print(f"Documento guardado en: {output_path}")

# =============================================================================
# BLOQUE PRINCIPAL
# =============================================================================

if __name__ == "__main__":
    # Ejemplo de estructura JSON con tres niveles
    json_data = {
        "SEGMENTOS": {
            "Segmentos Tradicional": [
                "Segemntación Demográfica",
                "Segmentación Psicografica",
                "Segmentacion Comportamental"
            ],
            "Segmentos Basados en el Concepto de JobstobeDone": [
                "Trabajos Funcionales",
                "Trabajos Sociales",
                "Trabajos Emocionales"                
            ],

            "Roles del Segmento": [
                "Clientes Directos vs. Clientes Indirectos (Intermediarios)",
                "Mercados de Doble Cara",
                "Segmentos en Mercados B2B (BusinesstoBusiness)",
                "Usuarios Gratuitos como Segmento de Clientes"
            ]
        },
        "PROPUESTA DE VALOR": {
            "Tipos de Propuesta de Valor": [
                "Propuesta de Valor Funcional",
                "Propuesta de Valor Social",
                "Propuesta de Valor Emocional"
            ],
            "Ejemplos de Aspectos Clave en la Propuesta de valor": [
                "Novedad e Innovación",
                "Mejora del Rendimiento",
                "Personalización y Diseño",
                "Marcad y estatus",
                "Reducción de Costes y Riesgos",
                "Accesibilidad y Comodidad"
            ]
        },
        "CANALES": {
            "Canales Directos e Indirectos": [
                "Canales directos",
                "Canales Indirectos"
            ],
             "Estrategia de Canales": [
                "Combinación de Canales",
                "Fases de los Canales y Roles"
            ]
        },
        "RELACION CON LOS CLIENTES": {
            "Tipos de Relaciones": [
                "Relación Directa",
                "Relación indirecta",
                "Relación Transaccional",
                "Relación a Largo Plazo",
                "Relación Automatizada",
                "Relación Personalizada"
            ],
            "Ciclo de Vida del Cliente": [
                "Adquisición, Retención y Expansión"
            ]
        },
        "INGRESOS": {
            "Flujos de Ingresos": [
                "Venta de Activos",
                "Cuotas de Uso",
                "Suscripción",
                "Préstamo",
                "Alquiler",
                "Leasing",
                "Cuotas de licencia",
                "Cuotas de publicidad"
            ],
            "Mecanismos de Fijación de Precios": [
                "Precios Estáticos vs. Dinámicos"
            ]
        },
        "ACTIVIDADES CLAVE": {
            "Configuraciones de Actividades": [
                "Cadenas de Producción",
                "Resolución de Problemas",
                "Gestión de Plataformas o Redes"
            ],
            "Consideraciones Estratégicas": [
                "Advertencia sobre el Granularismo y Consejo Estratégico"
            ]
        },
        "RECURSOS CLAVE": {
            "Recursos Tangibles e Intangibles": [
                "Recursos Físicos, Humanos y Financieros",
                "Propiedad Intelectual, Marca y Confianza"
            ],
            "Consideraciones Estratégicas": [
                "Selección de Recursos Críticos y Ejemplo de Decisión Estratégica"
            ]
        },
        "SOCIOS CLAVE": {
            "Tipos y Formas de Alianzas": [
                "Alianzas Estratégicas con NO competidores",
                "Cooperaci'on con competidores (Coopetición)",
                "Joint Ventures",
                "Relaciones Estrechas entre comprador y proveedor",
                "Adquisición de recursos o activiaddes"
            ],
            "Optimización y Reducción de Riesgos": [
                "Economías de Escala y Adquisición de Recursos",
                "Consideraciones Estratégicas para Socios Clave"
            ]
        },
        "COSTOS": {
            "Estructura de Costos": [
                "Costos Fijos y Variables",
                "Transición de Costos fijo a costo variable",
                "Modelos Basados en Costos"
            ],
            "Estrategias de Costos": [
                "Economías de Escala, de Alcance y Modelos Basados en Valor"
            ]
        }
    }

    # Configuración de rutas
    output_path = r"C:\Users\HP\Desktop\LIBROS PERSO\MODELOS DE NEGOCIOS\libro_BMC.docx"
    ruta_estilo = r"C:\Users\HP\Desktop\LIBROS PERSO\CONTEXTO ESPANIOL\CONTEXTO5.pdf"
    ruta_respuestas = r"C:\Users\HP\Desktop\2025-1-CATO-CURSOSO\GER-TI CATO-1-2025\RECURSOS\TEXTOS CANVAS\TEXTO DE VIDEOS\TEXTO VIDEOS EN ESPAÑOL\BMC TOTAL.pdf"

    # Generar el documento Word
    generate_word_document(json_data, output_path, ruta_estilo, ruta_respuestas)


# In[ ]:





# In[ ]:





# In[ ]:





# In[ ]:





# In[2]:


import os
from docx import Document
from dotenv import load_dotenv
from langchain.prompts import PromptTemplate
from langchain.vectorstores import FAISS
import nltk
from nltk.tokenize import sent_tokenize
from pydantic import BaseModel
from langchain_openai import ChatOpenAI, OpenAIEmbeddings

# Descargar recursos de NLTK (descomentar si es la primera vez)
# nltk.download('punkt')

# =============================================================================
# CONFIGURACIÓN Y CLIENTES
# =============================================================================

load_dotenv()
API_KEY = os.getenv("OPENAI_API_KEY")
if not API_KEY:
    raise ValueError("La clave de la API de OpenAI no se encontró en las variables de entorno.")

# Inicializar cliente LLM (se configura el modelo y otros parámetros según se requiera)
llm = ChatOpenAI(api_key=API_KEY, model="o3-mini-2025-01-31", temperature=1)

# Inicializar embeddings para vectorización
embeddings = OpenAIEmbeddings(api_key=API_KEY)

# =============================================================================
# FUNCIONES PARA EXTRAER Y PROCESAR TEXTO DE ARCHIVOS PDF
# =============================================================================

def get_pdf_text(ruta_pdf: str, max_tokens: int = 100000) -> str:
    """
    Lee y extrae el texto de un PDF, limitándolo a max_tokens tokens.
    """
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(ruta_pdf)
        tokens_totales = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                tokens_page = page_text.split()
                tokens_totales.extend(tokens_page)
                if len(tokens_totales) >= max_tokens:
                    tokens_totales = tokens_totales[:max_tokens]
                    break
        print(f"Total tokens leídos del PDF: {len(tokens_totales)}")
        return " ".join(tokens_totales)
    except Exception as e:
        print(f"Error al leer el PDF: {e}")
        return ""

def chunk_text_to_sentences(texto: str, min_words: int = 70) -> list:
    """
    Divide el texto en oraciones y filtra aquellas que tengan más de min_words palabras.
    """
    oraciones = sent_tokenize(texto, language='spanish')
    return [oracion.strip() for oracion in oraciones if len(oracion.split()) > min_words]

# =============================================================================
# CREACIÓN DE VECTORSTORES (EMBEDDINGS) BASADOS EN PDF
# =============================================================================

def create_style_vectorstore(ruta_estilo: str) -> FAISS:
    """
    Crea un vectorstore a partir de un archivo PDF de estilo.
    """
    if not ruta_estilo.lower().endswith('.pdf'):
        raise ValueError("El archivo de estilo debe estar en formato PDF.")

    texto_estilo = get_pdf_text(ruta_estilo, max_tokens=100000)
    oraciones = chunk_text_to_sentences(texto_estilo, min_words=70)
    if not oraciones:
        raise ValueError("No se encontraron oraciones válidas en el archivo de estilo.")
    return FAISS.from_texts(oraciones, embeddings)

def create_respuestas_vectorstore(ruta_respuestas: str) -> FAISS:
    """
    Crea un vectorstore a partir de un archivo PDF con frases de respuestas de otros libros.
    """
    if not ruta_respuestas.lower().endswith('.pdf'):
        raise ValueError("El archivo de respuestas debe estar en formato PDF.")

    texto_respuestas = get_pdf_text(ruta_respuestas, max_tokens=100000)
    oraciones = chunk_text_to_sentences(texto_respuestas, min_words=40)
    if not oraciones:
        raise ValueError("No se encontraron oraciones válidas en el archivo de respuestas.")
    return FAISS.from_texts(oraciones, embeddings)

def retrieve_style_text(vectorstore: FAISS, query: str = "Extrae las oraciones que mejor representen un estilo formal, académico y fluido", k: int = 10) -> str:
    """
    Recupera oraciones representativas del estilo deseado.
    """
    resultados = vectorstore.similarity_search(query, k=k)
    oraciones = [doc.page_content for doc in resultados]
    return "\n".join(oraciones)

def retrieve_respuestas_context(vectorstore: FAISS, query: str, k: int = 5) -> str:
    """
    Recupera frases relevantes del vectorstore de respuestas de otros libros.
    """
    resultados = vectorstore.similarity_search(query, k=k)
    frases = [doc.page_content for doc in resultados]
    return "\n".join(frases)

# =============================================================================
# TEMPLATES PARA GENERACIÓN DE RESPUESTAS
# =============================================================================

PROMPT_TITULO_EXPLICACION = PromptTemplate(
    input_variables=["titulo", "contexto_libros"],
    template="""Proporciona una explicación clara y detallada con dos ejemplos ilustrativos para el título '{titulo}' en el contexto del Business Model Canvas.
Utiliza el siguiente contexto extraído de fuentes académicas y otros libros:
{contexto_libros}
Asegúrate de que la respuesta incluya una exposición teórica detallada, profunda y académica, analice el concepto y su utilidad, y presente dos ejemplos prácticos en un escenario real.
La redacción debe ser clara, formal y académica, con oraciones extensas (al menos 50 palabras) y al menos tres oraciones concatenadas por un punto seguido.
Respuesta:"""
)

PROMPT_SUBTITULO_EXPLICACION = PromptTemplate(
    input_variables=["titulo", "subtitulo", "contexto_libros"],
    template="""Proporciona una explicación clara y detallada con dos ejemplos ilustrativos para el subtítulo '{subtitulo}', que forma parte del título '{titulo}' en el contexto del Business Model Canvas.
Utiliza el siguiente contexto extraído de fuentes académicas y otros libros:
{contexto_libros}
Asegúrate de que la respuesta incluya una exposición teórica detallada, profunda y académica, analice el concepto y su utilidad, y presente dos ejemplos prácticos en un escenario real.
La redacción debe ser clara, formal y académica, con oraciones extensas (al menos 50 palabras) y al menos tres oraciones concatenadas por un punto seguido.
Respuesta:"""
)

PROMPT_SUBSUBTITULO_EXPLICACION = PromptTemplate(
    input_variables=["titulo", "subtitulo", "subsubtitulo", "contexto_libros"],
    template="""Proporciona una explicación clara y detallada con dos ejemplos ilustrativos para el subtítulo de tercer nivel '{subsubtitulo}', que forma parte del subtítulo '{subtitulo}' del título '{titulo}' en el contexto del Business Model Canvas.
Utiliza el siguiente contexto extraído de fuentes académicas y otros libros:
{contexto_libros}
Asegúrate de que la respuesta incluya una exposición teórica detallada, profunda y académica, analice el concepto y su utilidad, y presente dos ejemplos prácticos en un escenario real.
La redacción debe ser clara, formal y académica, con oraciones extensas (al menos 50 palabras) y al menos tres oraciones concatenadas por un punto seguido.
Respuesta:"""
)

PROMPT_APPLY_HUMANO = PromptTemplate(
    input_variables=["contenido_original", "estilo"],
    template="""Ajusta tu forma de redacción para cumplir con los siguientes criterios, asegurando un texto que se asemeje a una respuesta completamente humana y alejada de los patrones de escritura típicos de inteligencia artificial.

Estilo de Redacción:
Imita la fluidez, la coherencia y la sofisticación de un texto académico extraído de un documento formal en español. Para ello, ten en cuenta las siguientes directrices:

Extensión y Complejidad de las Frases:
Redacta oraciones extensas que integren varias ideas de manera armónica, utilizando conectores fluidos en español.

Estructura Cohesiva y Fragmentación Adecuada:
Organiza el contenido en un mínimo de dos y un máximo de cinco oraciones interconectadas, asegurando una redacción natural y evitando frases inconexas.

Claridad y Consistencia:
Cada oración debe mantener un balance entre precisión y fluidez, evitando repeticiones innecesarias o explicaciones redundantes.

Uso Estratégico de Conectores:
Minimiza los puntos aislados y favorece una conexión fluida de ideas mediante conjunciones variadas, sin caer en estructuras mecánicas.

Diversificación Sintáctica:
Utiliza una combinación de oraciones subordinadas, construcciones con participios y elementos discursivos que den variedad y riqueza al texto.

Tono Académico pero Accesible:
Mantén un estilo formal y estructurado, pero sin tecnicismos excesivos que dificulten la comprensión. Piensa en la redacción de un libro universitario.

Extensión de las Frases:
Cada frase debe contener más de 50 palabras, integrando al menos tres proposiciones, asegurando un flujo de información profundo y detallado.

Instrucciones Finales:
Redacta el siguiente contenido aplicando todas las reglas anteriores, asegurando que el resultado sea indistinguible de un texto humano de alto nivel académico. Evita frases cortas, patrones predecibles o un tono excesivamente robótico. Usa la referencia proporcionada a continuación como base para emular el estilo lingüístico y discursivo.

Ejemplos de redacción: Utiliza de ejemplo de redacción las frases siguientes {estilo}

Reformula el siguiente contenido:

{contenido_original}
aplicándole el siguiente estilo humano deseado:
{estilo}
El resultado debe ser una redacción académica, fluida, con oraciones extensas (más de 50 palabras) y al menos tres oraciones, separadas por un punto seguido, manteniendo un tono cercano pero profesional.
Respuesta:"""
)

# =============================================================================
# FUNCIONES DE GENERACIÓN DE TEXTO
# =============================================================================

def generate_title_explanation(titulo: str, vectorstore_respuestas: FAISS, k: int = 10) -> str:
    """
    Genera la explicación y el ejemplo para un título, integrando contexto extraído mediante RAG de otros libros.
    """
    query = f"Explicación y ejemplo académico para el título '{titulo}' en el contexto del Business Model Canvas"
    contexto_libros = retrieve_respuestas_context(vectorstore_respuestas, query, k)
    formatted_prompt = PROMPT_TITULO_EXPLICACION.format(
        titulo=titulo,
        contexto_libros=contexto_libros
    )
    response = llm.invoke(formatted_prompt)
    return response.content.strip()

def generate_subtitle_explanation(titulo: str, subtitulo: str, vectorstore_respuestas: FAISS, k: int = 10) -> str:
    """
    Genera la explicación y el ejemplo para un subtítulo, integrando contexto extraído mediante RAG de otros libros.
    """
    query = f"Explicación y ejemplo académico para el subtítulo '{subtitulo}' del título '{titulo}' en el contexto del Business Model Canvas"
    contexto_libros = retrieve_respuestas_context(vectorstore_respuestas, query, k)
    formatted_prompt = PROMPT_SUBTITULO_EXPLICACION.format(
        titulo=titulo,
        subtitulo=subtitulo,
        contexto_libros=contexto_libros
    )
    response = llm.invoke(formatted_prompt)
    return response.content.strip()

def generate_subsubtitulo_explanation(titulo: str, subtitulo: str, subsubtitulo: str, vectorstore_respuestas: FAISS, k: int = 10) -> str:
    """
    Genera la explicación y el ejemplo para un subtítulo de tercer nivel,
    integrando contexto extraído mediante RAG de otros libros.
    """
    query = f"Explicación y ejemplo académico para '{subsubtitulo}' en el contexto de '{titulo}' > '{subtitulo}'"
    contexto_libros = retrieve_respuestas_context(vectorstore_respuestas, query, k)
    formatted_prompt = PROMPT_SUBSUBTITULO_EXPLICACION.format(
        titulo=titulo,
        subtitulo=subtitulo,
        subsubtitulo=subsubtitulo,
        contexto_libros=contexto_libros
    )
    response = llm.invoke(formatted_prompt)
    return response.content.strip()

def apply_human_style_to_text(contenido_original: str, estilo: str) -> str:
    """
    Transforma el contenido original aplicándole el estilo humano deseado.
    """
    formatted_prompt = PROMPT_APPLY_HUMANO.format(
        contenido_original=contenido_original,
        estilo=estilo
    )
    response = llm.invoke(formatted_prompt)
    return response.content.strip()

# =============================================================================
# GENERACIÓN DEL DOCUMENTO WORD A PARTIR DE UNA ESTRUCTURA JSON
# =============================================================================

def generate_word_document(json_data: dict, output_path: str, ruta_estilo: str, ruta_respuestas: str):
    """
    Procesa un JSON con la estructura jerárquica del Business Model Canvas y genera,
    para cada título, subtítulo y subtítulo de tercer nivel, una explicación académica detallada enriquecida
    con contexto de otros libros (usando RAG) y adaptada al estilo humano deseado.
    """
    # Crear vectorstores
    vectorstore_estilo = create_style_vectorstore(ruta_estilo)
    texto_estilo = retrieve_style_text(vectorstore_estilo)
    vectorstore_respuestas = create_respuestas_vectorstore(ruta_respuestas)

    doc = Document()
    doc.add_heading("Libro Business Model Canvas", level=1)

    for titulo, subtitulos in json_data.items():
        doc.add_heading(titulo, level=1)
        # Generar y agregar explicación para el título
        titulo_explicacion = generate_title_explanation(titulo, vectorstore_respuestas, k=10)
        titulo_explicacion_estilizada = apply_human_style_to_text(titulo_explicacion, texto_estilo)
        doc.add_paragraph(titulo_explicacion_estilizada)

        for subtitulo, subsubtitulos in subtitulos.items():
            doc.add_heading(subtitulo, level=2)
            # Generar y agregar explicación para el subtítulo
            subtitulo_explicacion = generate_subtitle_explanation(titulo, subtitulo, vectorstore_respuestas, k=10)
            subtitulo_explicacion_estilizada = apply_human_style_to_text(subtitulo_explicacion, texto_estilo)
            doc.add_paragraph(subtitulo_explicacion_estilizada)

            for subsubtitulo in subsubtitulos:
                print(f"Procesando: {titulo} > {subtitulo} > {subsubtitulo}")
                # Generar explicación para el tercer nivel
                subsubtitulo_explicacion = generate_subsubtitulo_explanation(titulo, subtitulo, subsubtitulo, vectorstore_respuestas, k=10)
                subsubtitulo_explicacion_estilizada = apply_human_style_to_text(subsubtitulo_explicacion, texto_estilo)
                doc.add_heading(subsubtitulo, level=3)
                doc.add_paragraph(subsubtitulo_explicacion_estilizada)

    doc.save(output_path)
    print(f"Documento guardado en: {output_path}")

# =============================================================================
# BLOQUE PRINCIPAL
# =============================================================================

if __name__ == "__main__":
    # Ejemplo de estructura JSON con tres niveles
    json_data = {
        "SEGMENTOS": {
            "Segmentos Tradicional": [
                "Segemntación Demográfica",
                "Segmentación Psicografica",
                "Segmentacion Comportamental"
            ],
            "Segmentos Basados en el Concepto de JobstobeDone": [
                "Trabajos Funcionales",
                "Trabajos Sociales",
                "Trabajos Emocionales"                
            ],
            "Roles del Segmento": [
                "Clientes Directos vs. Clientes Indirectos (Intermediarios)",
                "Mercados de Doble Cara",
                "Segmentos en Mercados B2B (BusinesstoBusiness)",
                "Usuarios Gratuitos como Segmento de Clientes"
            ]
        },
        "PROPUESTA DE VALOR": {
            "Tipos de Propuesta de Valor": [
                "Propuesta de Valor Funcional",
                "Propuesta de Valor Social",
                "Propuesta de Valor Emocional"
            ],
            "Ejemplos de Aspectos Clave en la Propuesta de valor": [
                "Novedad e Innovación",
                "Mejora del Rendimiento",
                "Personalización y Diseño",
                "Marcad y estatus",
                "Reducción de Costes y Riesgos",
                "Accesibilidad y Comodidad"
            ]
        },
        "CANALES": {
            "Canales Directos e Indirectos": [
                "Canales directos",
                "Canales Indirectos"
            ],
             "Estrategia de Canales": [
                "Combinación de Canales",
                "Fases de los Canales y Roles"
            ]
        },
        "RELACION CON LOS CLIENTES": {
            "Tipos de Relaciones": [
                "Relación Directa",
                "Relación indirecta",
                "Relación Transaccional",
                "Relación a Largo Plazo",
                "Relación Automatizada",
                "Relación Personalizada"
            ],
            "Ciclo de Vida del Cliente": [
                "Adquisición, Retención y Expansión"
            ]
        },
        "INGRESOS": {
            "Flujos de Ingresos": [
                "Venta de Activos",
                "Cuotas de Uso",
                "Suscripción",
                "Préstamo",
                "Alquiler",
                "Leasing",
                "Cuotas de licencia",
                "Cuotas de publicidad"
            ],
            "Mecanismos de Fijación de Precios": [
                "Precios Estáticos vs. Dinámicos"
            ]
        },
        "ACTIVIDADES CLAVE": {
            "Configuraciones de Actividades": [
                "Cadenas de Producción",
                "Resolución de Problemas",
                "Gestión de Plataformas o Redes"
            ],
            "Consideraciones Estratégicas": [
                "Advertencia sobre el Granularismo y Consejo Estratégico"
            ]
        },
        "RECURSOS CLAVE": {
            "Recursos Tangibles e Intangibles": [
                "Recursos Físicos, Humanos y Financieros",
                "Propiedad Intelectual, Marca y Confianza"
            ],
            "Consideraciones Estratégicas": [
                "Selección de Recursos Críticos y Ejemplo de Decisión Estratégica"
            ]
        },
        "SOCIOS CLAVE": {
            "Tipos y Formas de Alianzas": [
                "Alianzas Estratégicas con NO competidores",
                "Cooperaci'on con competidores (Coopetición)",
                "Joint Ventures",
                "Relaciones Estrechas entre comprador y proveedor",
                "Adquisición de recursos o activiaddes"
            ],
            "Optimización y Reducción de Riesgos": [
                "Economías de Escala y Adquisición de Recursos",
                "Consideraciones Estratégicas para Socios Clave"
            ]
        },
        "COSTOS": {
            "Estructura de Costos": [
                "Costos Fijos y Variables",
                "Transición de Costos fijo a costo variable",
                "Modelos Basados en Costos"
            ],
            "Estrategias de Costos": [
                "Economías de Escala, de Alcance y Modelos Basados en Valor"
            ]
        }
    }

    # Configuración de rutas
    output_path = r"C:\Users\HP\Desktop\LIBROS PERSO\MODELOS DE NEGOCIOS\libro_BMC.docx"
    ruta_estilo = r"C:\Users\HP\Desktop\LIBROS PERSO\CONTEXTO ESPANIOL\CONTEXTO5.pdf"
    ruta_respuestas = r"C:\Users\HP\Desktop\2025-1-CATO-CURSOSO\GER-TI CATO-1-2025\RECURSOS\TEXTOS CANVAS\TEXTO DE VIDEOS\TEXTO VIDEOS EN ESPAÑOL\BMC TOTAL.pdf"

    # Generar el documento Word
    generate_word_document(json_data, output_path, ruta_estilo, ruta_respuestas)


# In[ ]:




